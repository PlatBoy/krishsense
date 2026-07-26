from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "KrishiSense_Presentation_Guide.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1D2B24"
MUTED = "5F6B62"
LIGHT_FILL = "E8EEF5"
SOFT_GREEN = "EAF4EA"


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    total_width_dxa = 9360
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_width_dxa))

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(grid_col)


def set_run_font(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("KrishiSense Project Presentation Guide")
    set_run_font(run, 24, True, DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Simple notes for explaining the farmer soil analysis and market platform")
    set_run_font(run, 12, False, MUTED)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, 16, True, BLUE)


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    set_run_font(run, 13, True, BLUE)


def add_para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead:
        lead = p.add_run(f"{bold_lead}: ")
        set_run_font(lead, 11, True, INK)
    run = p.add_run(text)
    set_run_font(run, 11, False, INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, 11, False, INK)
    return p


def add_callout(doc, label, text, fill=SOFT_GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run(f"{label}: ")
    set_run_font(lead, 11, True, DARK_BLUE)
    body = p.add_run(text)
    set_run_font(body, 11, False, INK)
    doc.add_paragraph()


def add_feature_table(doc):
    rows = [
        ("Soil photo analysis", "Farmer uploads a photo and gets soil type, confidence, health score, and suggestions."),
        ("Soil identifier", "A quick photo-only feature for identifying a random soil sample."),
        ("Loans", "Farmer applies for a loan. Admin approves or rejects it."),
        ("Market", "Farmer buys seeds, fertiliser, equipment, or a tractor using the account balance."),
        ("Admin portal", "Admin can review reports, approve loans, ban farmers, remove farmers, and reset farmer passwords."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [2.0, 4.5])
    hdr = table.rows[0].cells
    hdr[0].text = "Feature"
    hdr[1].text = "Simple explanation"
    for cell in hdr:
        set_cell_fill(cell, LIGHT_FILL)
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, 10.5, True, DARK_BLUE)

    for feature, explanation in rows:
        cells = table.add_row().cells
        cells[0].text = feature
        cells[1].text = explanation
        for cell in cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run_font(run, 10.5, False, INK)


def add_flow_table(doc):
    rows = [
        ("1", "Login/Register", "JWT token keeps the user session secure."),
        ("2", "Upload photo/details", "Express receives form data and image upload."),
        ("3", "Store image", "Cloudinary stores the photo safely in the cloud."),
        ("4", "Analyze with Gemini", "Gemini reads the image and returns soil classification and advice."),
        ("5", "Save report", "MongoDB saves user data, report results, loans, and orders."),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [0.65, 1.85, 4.0])
    hdr = table.rows[0].cells
    hdr[0].text = "Step"
    hdr[1].text = "Action"
    hdr[2].text = "What happens"
    for cell in hdr:
        set_cell_fill(cell, LIGHT_FILL)
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, 10.5, True, DARK_BLUE)

    for number, action, detail in rows:
        cells = table.add_row().cells
        cells[0].text = number
        cells[1].text = action
        cells[2].text = detail
        for cell in cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run_font(run, 10.5, False, INK)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_simple_architecture(doc):
    add_bullet(doc, "Frontend: React and Vite show the farmer dashboard, admin dashboard, forms, history, theme switch, and market.")
    add_bullet(doc, "Backend: Express.js handles login, image upload, soil analysis requests, loans, market orders, and admin actions.")
    add_bullet(doc, "Database: MongoDB stores users, reports, loans, wallet balances, and market orders.")
    add_bullet(doc, "Image storage: Cloudinary keeps uploaded soil photos safe even after deployment.")
    add_bullet(doc, "AI service: Google Gemini is used through an API. I did not train a model from scratch.")
    add_bullet(doc, "Deployment: Render hosts the live app, and GitHub stores the source code.")


def add_script_lines(doc):
    lines = [
        "My project is called KrishiSense. It is made for farmers who want simple help with soil and crop decisions.",
        "The farmer can upload a soil photo and enter land details. The backend sends the image to Gemini and stores the report in MongoDB.",
        "There is also an admin portal. The admin can review reports, manage farmers, approve loans, and reset passwords.",
        "I added a market section. When a loan is approved, the farmer account balance increases, and the farmer can buy seeds, fertiliser, equipment, or tractor items.",
        "The repayment option is a built-in demo bank form. It records repayment against an approved loan and masks the account number.",
        "For deployment, I used GitHub for code and Render for hosting. Environment variables keep API keys and database credentials out of the code.",
    ]
    for line in lines:
        add_bullet(doc, line)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    add_title(doc)
    add_callout(
        doc,
        "One-line idea",
        "KrishiSense is a web app that helps farmers understand soil, manage reports, apply for loans, and use loan money in a simple market.",
    )
    add_h1(doc, "1. What The Project Does")
    add_para(
        doc,
        "The main goal is to make a useful farmer support website. A farmer logs in, uploads a soil photo, adds land details, and receives a simple soil report. The same website also includes loan applications, an admin portal, and a market for buying farm items.",
    )
    add_feature_table(doc)
    add_h2(doc, "Why it is useful")
    add_bullet(doc, "Farmers get fast soil guidance without needing to understand technical lab reports.")
    add_bullet(doc, "Reports are saved, so the farmer can compare old and new soil results.")
    add_bullet(doc, "Loans and market purchases are connected in one simple flow.")
    add_bullet(doc, "Admins can manage users and decisions from a separate dashboard.")

    doc.add_page_break()
    add_h1(doc, "2. How The Website Works")
    add_para(
        doc,
        "The website has two main sides: farmer side and admin side. Farmers can submit data and use the market. Admins can review reports, approve loans, ban or remove farmers, and reset passwords.",
    )
    add_flow_table(doc)
    add_h2(doc, "Main technology")
    add_simple_architecture(doc)
    add_callout(
        doc,
        "Simple way to explain",
        "I used APIs and cloud services. I did not build my own AI model. The project connects the frontend, backend, database, image storage, and AI API together.",
        LIGHT_FILL,
    )

    doc.add_page_break()
    add_h1(doc, "3. Loan, Balance, Market, And Repayment")
    add_para(
        doc,
        "The new market feature makes the project feel more complete. The farmer can see an account balance, buy farm products, and repay approved loans using a demo bank form.",
    )
    add_h2(doc, "Loan to balance flow")
    add_bullet(doc, "The farmer fills a loan form with amount, purpose, crop, land area, and tenure.")
    add_bullet(doc, "The admin opens the Loans tab and approves or rejects the request.")
    add_bullet(doc, "If approved, the backend adds that loan amount to the farmer wallet balance one time.")
    add_bullet(doc, "The farmer can then buy seeds, fertiliser, equipment, or tractor items from the market.")
    add_bullet(doc, "For repayment, the farmer selects the loan, enters amount and bank details, and the website records the payment.")
    add_h2(doc, "Market items")
    add_bullet(doc, "Seeds: wheat seed, paddy seed, and vegetable seed kit.")
    add_bullet(doc, "Fertiliser: neem coated urea, DAP, and NPK.")
    add_bullet(doc, "Equipment: battery sprayer and drip irrigation kit.")
    add_bullet(doc, "Tractor: 45 HP tractor with an indicative ex-showroom price.")
    add_para(
        doc,
        "The prices are stored as a demo catalog in the backend. A real business version can replace this with a supplier API so prices update automatically.",
        "Important note",
    )

    doc.add_page_break()
    add_h1(doc, "4. What I Should Say In The Presentation")
    add_para(
        doc,
        "Keep the explanation simple. Do not say you trained an AI model or built every cloud service yourself. Say that you built a full-stack web app and integrated existing services using APIs.",
    )
    add_h2(doc, "Short speaking script")
    add_script_lines(doc)
    add_h2(doc, "Questions the teacher may ask")
    add_para(doc, "What is JWT?", "Answer")
    add_para(doc, "JWT is a login token. After login, the frontend sends the token with requests so the backend knows which user is using the app.")
    add_para(doc, "Why MongoDB?", "Answer")
    add_para(doc, "MongoDB stores flexible data like user profiles, soil reports, loan records, market orders, and repayment history.")
    add_para(doc, "How does soil analysis work?", "Answer")
    add_para(doc, "The app sends the uploaded soil photo and farmer details to Gemini. Gemini returns soil type, confidence, and recommendations. The result is saved in the database.")
    add_para(doc, "How did you deploy it?", "Answer")
    add_para(doc, "I pushed the code to GitHub and deployed the Node/React app on Render. Secret keys are set as environment variables.")
    add_callout(
        doc,
        "Safe final line",
        "This project is not meant to replace agricultural experts. It is a student prototype that shows how technology can support farmers with faster information.",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("KrishiSense presentation guide")
    set_run_font(run, 9, False, MUTED)

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
